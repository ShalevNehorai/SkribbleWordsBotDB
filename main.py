import os
import re
import discord
from discord.ext import commands
from dotenv import load_dotenv
from keep_alive import keep_alive
import database_helper as helper

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

TOKEN = os.getenv('DISCORD_TOKEN')

bot = commands.Bot(command_prefix='!')
bot.remove_command('help')

@bot.event
async def on_ready():
  print(f'{bot.user.name} has connected to Discord!')
  await bot.change_presence(activity=discord.Game('skribbl.io'))

# @bot.command(name="play", help='send link to Skribbl.io')
# async def play(ctx):
#   await ctx.send("https://skribbl.io/")

async def send_words_file(ctx, words):
  if words:
    with open("words.txt", "w") as file:
      file.write(','.join(words))
    with open("words.txt", "rb") as file:
      await ctx.send("Your file is:", file=discord.File(file, "words.txt"))
  else:
    await ctx.send('no words has been found')

@bot.command(name="all", help="return file with all the words in the DB")
async def get_all_words(ctx):
  await ctx.send('making the file')
  words = helper.get_all_words()
  await send_words_file(ctx, words)

@bot.command(name="indexed", help="return indexed file of all the words")
async def get_all_words_sorted(ctx):
  await ctx.send('making the file')
  words = helper.get_all_words()
  if words:
    with open("words.txt", "w") as file:
      file.write("\n".join([f'{i+1}: {w}' for i, w in enumerate(words)]))
    with open("words.txt", "rb") as file:
      await ctx.send("Your file is:", file=discord.File(file, "words.txt"))
  else:
    await ctx.send('no words has been found')

@bot.command(name="first", help='argu: number:int return file with the {number} first words')
async def get_first_words(ctx, number:int):
  await ctx.send('making the file')
  words = helper.get_first_words(number)
  await send_words_file(ctx, words)

@bot.command(name="last", help='argu: number:int return file with the {number} last words')
async def get_last_words(ctx, number:int):
  await ctx.send('making the file')
  words = helper.get_last_words(number)
  await send_words_file(ctx, words)

@bot.command(name="random", help='argu: number:int last(optional); return file with {number} random words and {last} last words')
async def get_random_words(ctx, number:int, last = 0):
  await ctx.send('making the file')
  words = helper.get_random_words(number, last)
  await send_words_file(ctx, words)

@bot.command(name="words-in-chat", help="write 75 words in chat")
async def get_words_in_chat(ctx):
  await ctx.send('working on it')
  words = helper.get_random_words(75)
  if words:
    await ctx.send(words)
  else:
    await ctx.send('no words has been found')

@bot.command(name="first-last", help="return file with {number} first and last words")
async def get_first_last(ctx, number:int):
  await ctx.send('making the file')
  words = helper.get_first_last(number)
  await send_words_file(ctx, words)

@bot.command(name="count-all", help='print the amount of words in the DB')
async def get_words_number(ctx):
  num = helper.count_words()
  await ctx.send(f'we have **{num}** words!')

@bot.command(name="count-new", help='print the amount of words that added in the past 48 hours')
async def count_new_words(ctx):
  num = helper.count_new_words()
  await ctx.send(f'**{num}** words added')

# @bot.command(name="count-my" )
# async def count_my_words(ctx):
#   author = ctx.message.author
#   num = helper.count_by_author(author.name)
#   await ctx.send(f'{author.name} you wrote **{num}** words')

@bot.command(name='author', help='argu: word:str print the word author')
async def get_author(ctx, *, word):
  author = helper.get_author(word)
  if author is not None:
    await ctx.send(f'**{author}** is the author')
  else:
    await ctx.send("Word not found")

@bot.command(name='stats-all', help="print the number of words each author added to the DB")
async def get_stats(ctx):
  await write_stats(ctx, helper.stats_all())

@bot.command(name='stats-new', help="print the number of words each author added to the DB in the last 48 hours")
async def get_stats_new(ctx):
  await write_stats(ctx, helper.stats_new())

async def write_stats(ctx, stats):
  msg = ""
  for stat in stats:
    msg += f'{stat[0]} wrote **{stat[1]}** words\n'
  await ctx.send(msg)

@bot.command(name='add', help='argu: arg:str add the coma separeted words in {arg} to the DB')
async def add_words(ctx, *, arg):
  MAX_WORD_LENGTH = 30
  word_list = arg.split(',')
  author = ctx.message.author
  msg_date = ctx.message.created_at
  count_words = 0
  for word in word_list:
    word = word.replace('\'', '')
    if(len(word) > MAX_WORD_LENGTH):
      await ctx.send(f'{word} it is to long, maximum characters allowed is ' + str(MAX_WORD_LENGTH) + " chracters")
    elif bool(re.match('[א-ת0-9\s?!]+$', word)):
      helper.__add_word__(word, author.name, msg_date)
      count_words += 1
    else:
      await ctx.send(f'{word} containe non aturaize characters and not added to the list')

  await ctx.send(f'{author.name} has added {count_words} word to the list')
  try:
    await ctx.message.delete()
  except:
    pass

@bot.command(name='CODENAME', help="send docx file with table of random words")
async def CODENAME(ctx):
  words = helper.get_random_words(25)
  document = Document()
  table = document.add_table(rows=5, cols=5)
  table.style = document.styles['Table Grid']
  i = 0
  for row in table.rows:
    for cell in row.cells:
      cell.text = words[i]
      i += 1
      cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
      cell.paragraphs[0].runs[0].font.size = Pt(20)
  document.save("table.docx")
  with open('table.docx', 'rb') as file:
    await ctx.send("Your file is:", file=discord.File(file, "words.docx"))


@bot.command(name="help", help="show this message")
async def help(ctx, command=None):
  if(command is None):
    embed = discord.Embed(
      title="help", colour= discord.Colour.blue()
    )
    embed.set_footer(text="Type !help command for more info on a command.")

    for com in bot.commands:
      # print(com.signature)
      embed.add_field(name=com.name, value=com.help, inline=False)

    await ctx.send(embed = embed)
  else:
    com = bot.get_command(command)
    if com is not None:
      embed = discord.Embed(
        title=com.name, colour= discord.Colour.blue()
      )
      params = ""
      for v in com.clean_params.items():
       params += str(v[1]) + ", "
      if params == "":
        params = "no parameters"
      embed.add_field(name=params, value=com.help)

      await ctx.send(embed = embed)
    else:
      await ctx.send("command not found")

@bot.event
async def on_command_error(ctx, error):
  if isinstance(error, commands.errors.CheckFailure):
    await ctx.send('you don\'t have permission for this command')
  elif isinstance(error, commands.errors.MissingRequiredArgument):
    await ctx.send(f'missing command argument: {error.param}')
  elif isinstance(error, commands.errors.BadArgument):
    await ctx.send('parameter must be number')
  elif isinstance(error, commands.errors.CommandNotFound):
    print('command not found')
    await ctx.send('command not found')
  else:
    await ctx.send('some error occured with the command')
    guild = ctx.guild
    log_channel = discord.utils.get(guild.channels, name='error-logs')
    if log_channel:
      await log_channel.send(error)
    else:
      await ctx.send(error)

keep_alive()
bot.run(TOKEN)